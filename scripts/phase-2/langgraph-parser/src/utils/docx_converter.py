"""
DOCX to Markdown 변환 유틸리티 (python-docx 기반)

명세서 4.0 참조:
- Word 하이라이트 → [text]{.mark}
- Word 볼드/이탤릭/밑줄 → Markdown 서식
- TOC 스타일에서 직접 TOC 구조 추출
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def extract_meeting_id(file_path: str) -> str:
    """파일명에서 meeting_id 추출

    패턴:
        - Final_Minutes_report_RAN1#120_v100.docx → RAN1_120
        - TSGR1_118_FinalMinutes.docx → RAN1_118

    Args:
        file_path: .docx 파일 경로

    Returns:
        meeting_id (예: RAN1_120)
    """
    filename = Path(file_path).stem

    # 패턴 1: RAN1#120
    match = re.search(r'RAN(\d+)#(\d+)', filename)
    if match:
        return f"RAN{match.group(1)}_{match.group(2)}"

    # 패턴 2: TSGR1_120
    match = re.search(r'TSGR1_(\d+)', filename)
    if match:
        return f"RAN1_{match.group(1)}"

    return "unknown"


def process_run(run) -> str:
    """개별 Run의 서식을 Markdown으로 변환

    서식 중첩 처리 (안쪽에서 바깥쪽으로):
        1. 밑줄 → <u>text</u>
        2. 이탤릭 → *text*
        3. 볼드 → **text**
        4. 하이라이트 → [text]{.mark}

    Args:
        run: docx.text.run.Run 객체

    Returns:
        서식이 적용된 Markdown 텍스트
    """
    text = run.text
    if not text or not text.strip():
        return text if text else ""

    # 취소선 확인
    has_strike = False
    if run._r.rPr is not None:
        strike_elem = run._r.rPr.find(qn('w:strike'))
        if strike_elem is not None:
            has_strike = strike_elem.get(qn('w:val'), 'true') != 'false'

    # 서식 적용 (안쪽에서 바깥쪽으로)
    if has_strike:
        text = f"~~{text}~~"
    if run.underline:
        text = f"<u>{text}</u>"
    if run.italic:
        text = f"*{text}*"
    if run.bold:
        text = f"**{text}**"

    # 하이라이트 확인 (최외곽) - Spec 4.0: 색상별 마커 생성
    if run.font.highlight_color is not None:
        from docx.enum.text import WD_COLOR_INDEX
        color = run.font.highlight_color
        # 색상별 마커 매핑 (Spec Line 1282-1285)
        if color == WD_COLOR_INDEX.BRIGHT_GREEN:
            text = f"[{text}]{{.mark-green}}"
        elif color == WD_COLOR_INDEX.YELLOW or color == WD_COLOR_INDEX.DARK_YELLOW:
            text = f"[{text}]{{.mark-yellow}}"
        elif color == WD_COLOR_INDEX.TURQUOISE:
            text = f"[{text}]{{.mark-turquoise}}"
        else:
            text = f"[{text}]{{.mark}}"

    return text


def get_paragraph_style_name(para) -> str:
    """Paragraph의 스타일 이름 반환

    Args:
        para: docx.text.paragraph.Paragraph 객체

    Returns:
        스타일 이름 (예: "TOC 1", "Heading 1", "Normal")
    """
    if para.style is not None:
        return para.style.name
    return "Normal"


def is_toc_style(style_name: str) -> bool:
    """TOC 스타일인지 확인

    Args:
        style_name: 스타일 이름

    Returns:
        TOC 스타일이면 True
    """
    # 대소문자 무시: "TOC 1", "toc 1" 모두 처리
    return style_name.lower().startswith("toc ")


def get_toc_depth(style_name: str) -> int:
    """TOC 스타일에서 depth 추출

    Args:
        style_name: 스타일 이름 (예: "TOC 1", "toc 1")

    Returns:
        depth (1, 2, 3, ...)
    """
    # 대소문자 무시
    match = re.match(r'toc\s+(\d+)', style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return 1


def extract_page_number(text: str) -> Optional[int]:
    """TOC 항목 텍스트에서 페이지 번호 추출

    Args:
        text: TOC 항목 텍스트 (예: "1 Opening of the meeting 5")

    Returns:
        페이지 번호 (없으면 None)
    """
    page_match = re.search(r'\s+(\d+)\s*$', text)
    return int(page_match.group(1)) if page_match else None


def has_section_number(text: str) -> bool:
    """텍스트가 섹션 번호로 시작하는지 확인

    Args:
        text: TOC 항목 텍스트

    Returns:
        섹션 번호가 있으면 True
    """
    # 숫자로 시작하는 패턴: 1, 1.1, 9.1.1 등
    return bool(re.match(r'^\d+(?:\.\d+)*\s+', text))


def generate_anchor(text: str) -> str:
    """TOC 항목 텍스트에서 anchor 생성

    Spec 4.0: anchor는 hyperlink에서 추출하거나 제목에서 생성

    Args:
        text: TOC 항목 텍스트

    Returns:
        anchor 문자열 (예: "opening-of-the-meeting")
    """
    # 페이지 번호 제거
    page_match = re.search(r'\s+\d+\s*$', text)
    if page_match:
        text = text[:page_match.start()].strip()

    # 섹션 번호 제거
    id_match = re.match(r'^(?:\d+(?:\.\d+)*)\s+', text)
    if id_match:
        text = text[id_match.end():].strip()

    # Annex 형식 처리: "Annex A-1: List of..." → "annex-a-1-list-of"
    # 또는 그대로 anchor로 변환

    # anchor 생성: 소문자, 공백→하이픈, 특수문자 제거
    anchor = text.lower()
    anchor = re.sub(r'[^\w\s-]', '', anchor)  # 특수문자 제거
    anchor = re.sub(r'\s+', '-', anchor)  # 공백 → 하이픈
    anchor = re.sub(r'-+', '-', anchor)  # 중복 하이픈 제거
    anchor = anchor.strip('-')

    return anchor


def extract_anchor_from_hyperlink(para) -> Optional[str]:
    """TOC 항목의 하이퍼링크에서 anchor 추출

    Word TOC는 하이퍼링크로 본문 헤더에 연결됨.
    이 함수는 하이퍼링크의 anchor를 추출.

    Args:
        para: docx.text.paragraph.Paragraph 객체

    Returns:
        anchor 문자열 (없으면 None)
    """
    # python-docx에서 하이퍼링크 추출
    # TOC 항목은 보통 w:hyperlink 요소로 감싸져 있음
    try:
        for hyperlink in para._element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            anchor = hyperlink.get(qn('w:anchor'))
            if anchor:
                return anchor
    except Exception:
        pass

    return None


def parse_toc_entry(text: str) -> tuple[Optional[str], str, Optional[int]]:
    """TOC 항목 텍스트 파싱 (레거시 호환용)

    형식: "9.1.1 MIMO 14" 또는 "MIMO 14" (번호 없음)

    Args:
        text: TOC 항목 텍스트

    Returns:
        tuple: (section_id, title, page)
            - section_id: 섹션 번호 (없으면 None)
            - title: 섹션 제목
            - page: 페이지 번호 (없으면 None)
    """
    text = text.strip()
    if not text:
        return None, "", None

    # 페이지 번호 추출 (마지막 숫자)
    page_match = re.search(r'\s+(\d+)\s*$', text)
    page = int(page_match.group(1)) if page_match else None
    if page_match:
        text = text[:page_match.start()].strip()

    # 섹션 번호 추출 (시작 부분)
    # 패턴: 9, 9.1, 9.1.1, Annex A, Annex B-1, etc.
    id_match = re.match(r'^((?:\d+(?:\.\d+)*)|(?:Annex\s+[A-Z](?:-\d+)?))\s+', text)
    if id_match:
        section_id = id_match.group(1).strip()
        title = text[id_match.end():].strip()
        return section_id, title, page

    # 번호 없는 경우
    return None, text, page


def extract_toc_from_docx(doc: Document) -> dict:
    """Word 문서에서 TOC 구조 직접 추출 (Spec 4.0 준수)

    TOC 스타일(TOC 1, TOC 2, TOC 3)을 사용하여 구조 추출

    Spec 4.0 toc_raw.yaml 형식:
        entries:
          - text: "1 Opening of the meeting"
            style: "TOC 1"
            depth: 1
            page: 5
            anchor: "opening-of-the-meeting"
          - text: "MIMO"
            style: "TOC 3"
            depth: 3
            page: 14
            anchor: "mimo"
            unnumbered: true

    Args:
        doc: Document 객체

    Returns:
        dict: Spec 4.0 toc_raw.yaml 구조
            - entries: TOC 항목 리스트
    """
    entries = []

    for para in doc.paragraphs:
        style_name = get_paragraph_style_name(para)

        if is_toc_style(style_name):
            # TOC 텍스트 추출 (하이퍼링크 포함)
            text = para.text.strip()
            if not text:
                continue

            depth = get_toc_depth(style_name)
            page = extract_page_number(text)

            # anchor 추출: 하이퍼링크에서 먼저 시도, 없으면 텍스트에서 생성
            anchor = extract_anchor_from_hyperlink(para)
            if not anchor:
                anchor = generate_anchor(text)

            # Spec 4.0 구조
            entry = {
                "text": text,
                "style": style_name,
                "depth": depth,
                "page": page,
                "anchor": anchor,
            }

            # 번호 없는 항목에 unnumbered 플래그 추가
            if not has_section_number(text):
                entry["unnumbered"] = True

            entries.append(entry)

    return {"entries": entries}


def get_heading_level(para) -> Optional[int]:
    """Paragraph의 Heading 레벨 반환

    Args:
        para: Paragraph 객체

    Returns:
        Heading 레벨 (1-9) 또는 None
    """
    style_name = get_paragraph_style_name(para)

    # Heading 1, Heading 2, ... 패턴
    match = re.match(r'Heading\s+(\d+)', style_name)
    if match:
        return int(match.group(1))

    return None


def extract_paragraph_content(para) -> str:
    """Paragraph에서 모든 텍스트 추출 (hyperlink 포함)

    DOCX 구조:
    - para.runs: 일반 텍스트 Run들
    - w:hyperlink: 하이퍼링크 요소 (para.runs에 포함되지 않음!)

    Word 문서의 TDoc 번호(R1-XXXXXXX)는 대부분 hyperlink로 되어 있어서
    para.runs만 처리하면 TDoc 번호가 손실됨.

    이 함수는 paragraph의 XML을 직접 순회하여 모든 텍스트를 추출함.

    Args:
        para: Paragraph 객체

    Returns:
        서식이 적용된 전체 텍스트
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    result_parts = []
    p_element = para._element

    # paragraph의 모든 자식 요소를 순서대로 처리
    for child in p_element:
        tag = child.tag

        # 일반 Run (w:r)
        if tag == qn('w:r'):
            # Run 객체 생성하여 process_run 호출
            run = Run(child, para)
            result_parts.append(process_run(run))

        # Hyperlink (w:hyperlink) - TDoc 번호가 여기에 포함됨!
        elif tag == qn('w:hyperlink'):
            # hyperlink 내부의 모든 run을 처리
            for run_elem in child.findall(qn('w:r')):
                run = Run(run_elem, para)
                result_parts.append(process_run(run))

    return "".join(result_parts)


def convert_paragraph_to_markdown(para) -> str:
    """Paragraph을 Markdown으로 변환 (hyperlink 포함)

    Args:
        para: Paragraph 객체

    Returns:
        Markdown 텍스트
    """
    # Heading 처리
    heading_level = get_heading_level(para)
    if heading_level:
        text = extract_paragraph_content(para)
        return f"{'#' * heading_level} {text}"

    # 일반 텍스트 (hyperlink 포함)
    return extract_paragraph_content(para)


def convert_docx_to_markdown(
    input_path: str,
    output_dir: str,
) -> tuple[str, dict]:
    """DOCX 파일을 Markdown + TOC 구조로 변환 (Spec 4.0 준수)

    Args:
        input_path: .docx 파일 경로
        output_dir: 출력 디렉토리

    Returns:
        tuple[str, dict]: (markdown_content, toc_raw)
            - markdown_content: 서식 마커 포함 본문
            - toc_raw: Spec 4.0 toc_raw.yaml 구조 (entries 포함)

    Raises:
        FileNotFoundError: 파일이 없을 때
        ImportError: python-docx가 설치되지 않았을 때
    """
    input_path = Path(input_path)
    output_dir = Path(output_dir)

    if not input_path.exists():
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {input_path}")

    # 출력 디렉토리 생성
    output_dir.mkdir(parents=True, exist_ok=True)

    # 문서 로드
    doc = Document(str(input_path))

    # TOC 추출
    toc_raw = extract_toc_from_docx(doc)

    # 본문을 Markdown으로 변환
    markdown_lines = []
    in_toc = False

    for para in doc.paragraphs:
        style_name = get_paragraph_style_name(para)

        # TOC 영역 건너뛰기
        if is_toc_style(style_name):
            in_toc = True
            continue

        # TOC 직후 본문 시작 감지
        if in_toc and not is_toc_style(style_name):
            # TOC 영역이 끝났으면 본문 시작
            if para.text.strip():
                in_toc = False

        if in_toc:
            continue

        # Paragraph 변환
        md_text = convert_paragraph_to_markdown(para)
        if md_text:
            markdown_lines.append(md_text)
        else:
            # 빈 줄 유지
            markdown_lines.append("")

    markdown_content = "\n".join(markdown_lines)

    # 파일 저장
    import yaml

    # document.md 저장
    md_path = output_dir / "document.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(markdown_content)

    # toc_raw.yaml 저장
    toc_path = output_dir / "toc_raw.yaml"
    with open(toc_path, "w", encoding="utf-8") as f:
        yaml.dump(toc_raw, f, allow_unicode=True, default_flow_style=False)

    return markdown_content, toc_raw


if __name__ == "__main__":
    # 테스트
    import sys

    if len(sys.argv) < 3:
        print("Usage: python docx_converter.py <input.docx> <output_dir>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_dir = sys.argv[2]

    try:
        md_content, toc_raw = convert_docx_to_markdown(input_path, output_dir)
        entries = toc_raw.get("entries", [])
        print(f"✅ 변환 완료")
        print(f"   Markdown: {len(md_content):,} chars")
        print(f"   TOC entries: {len(entries)}")

        # 하이라이트 마커 확인
        mark_count = md_content.count("{.mark}")
        print(f"   Highlight markers: {mark_count}")

    except Exception as e:
        print(f"❌ 오류: {e}")
        sys.exit(1)
