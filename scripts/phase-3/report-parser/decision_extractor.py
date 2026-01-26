"""
Decision Extractor - Extract Agreements, Conclusions, Working Assumptions

Uses regex to detect decision keywords and LLM for boundary/validation.
Integrates with AgendaMatcher for correct agenda assignment.
"""

import re
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

from docx import Document

from models import Decision, DecisionType, Section


# Decision keyword patterns - Support both header-only and inline content
# Pattern 1: Header only (e.g., "Agreement:")
# Pattern 2: Inline content (e.g., "Agreement: The UE shall...")
DECISION_PATTERNS_HEADER = {
    DecisionType.AGREEMENT: re.compile(
        r'^agreements?\s*:?\s*$',
        re.IGNORECASE
    ),
    DecisionType.CONCLUSION: re.compile(
        r'^conclusions?\s*:?\s*$',
        re.IGNORECASE
    ),
    DecisionType.WORKING_ASSUMPTION: re.compile(
        r'^working\s*assumptions?\s*:?\s*$',
        re.IGNORECASE
    ),
}

# Inline patterns - match "Agreement: [content]" where content exists
DECISION_PATTERNS_INLINE = {
    DecisionType.AGREEMENT: re.compile(
        r'^agreements?\s*:\s*(.+)',
        re.IGNORECASE | re.DOTALL
    ),
    DecisionType.CONCLUSION: re.compile(
        r'^conclusions?\s*:\s*(.+)',
        re.IGNORECASE | re.DOTALL
    ),
    DecisionType.WORKING_ASSUMPTION: re.compile(
        r'^working\s*assumptions?\s*:\s*(.+)',
        re.IGNORECASE | re.DOTALL
    ),
}

# Annex section pattern - skip these sections (Spec requirement)
ANNEX_PATTERN = re.compile(r'^Annex\s+[A-Z]', re.IGNORECASE)

# Pattern to extract Tdoc references
TDOC_PATTERN = re.compile(r'R\d-\d{7}')

# Pattern to detect FFS (For Further Study)
FFS_PATTERN = re.compile(r'\bFFS\b', re.IGNORECASE)

# Pattern to detect TBD (To Be Determined)
TBD_PATTERN = re.compile(r'\bTBD\b', re.IGNORECASE)


@dataclass
class DecisionCandidate:
    """Candidate decision found by regex."""
    decision_type: DecisionType
    keyword_index: int  # Paragraph index of the keyword
    keyword_text: str  # The keyword text
    inline_content: Optional[str] = None  # Content on the same line as keyword


class DecisionExtractor:
    """
    Extract decisions from document paragraphs.

    Two-phase approach:
    1. Regex detection: Find decision keywords (high recall)
    2. Content extraction: Get content until next decision/heading
    """

    def __init__(self, sections: list[Section]):
        """
        Initialize with document sections.

        Args:
            sections: List of sections with agenda numbers
        """
        self.sections = sections
        self._section_by_index: dict[int, Section] = {}
        self._build_section_index()

    def _build_section_index(self) -> None:
        """Build index for quick section lookup by paragraph index."""
        for section in self.sections:
            for idx in range(section.start_index, section.end_index):
                self._section_by_index[idx] = section

    def extract_decisions(
        self,
        paragraphs: list,
        meeting_id: str
    ) -> tuple[list[Decision], list[Decision], list[Decision]]:
        """
        Extract all decisions from document paragraphs.

        Args:
            paragraphs: List of docx paragraphs
            meeting_id: Meeting ID for decision IDs

        Returns:
            Tuple of (agreements, conclusions, working_assumptions)
        """
        # Phase 1: Find candidates
        candidates = self._find_candidates(paragraphs)

        # Phase 2: Extract content for each candidate
        agreements = []
        conclusions = []
        working_assumptions = []

        # Track decision counts per agenda for ID generation
        counters: dict[tuple[DecisionType, str], int] = {}

        for i, candidate in enumerate(candidates):
            # Handle inline content vs header-only
            if candidate.inline_content:
                # Inline: content is on the same line
                # But may continue to following paragraphs
                next_boundary = self._find_content_boundary(
                    candidate.keyword_index,
                    paragraphs,
                    candidates[i + 1].keyword_index if i + 1 < len(candidates) else len(paragraphs)
                )
                following_content = self._extract_content(
                    paragraphs,
                    candidate.keyword_index + 1,
                    next_boundary
                )
                # Combine inline content with following content
                if following_content.strip():
                    content = candidate.inline_content + '\n' + following_content
                else:
                    content = candidate.inline_content
            else:
                # Header-only: content in following paragraphs
                next_boundary = self._find_content_boundary(
                    candidate.keyword_index,
                    paragraphs,
                    candidates[i + 1].keyword_index if i + 1 < len(candidates) else len(paragraphs)
                )
                content = self._extract_content(
                    paragraphs,
                    candidate.keyword_index + 1,  # Skip keyword paragraph
                    next_boundary
                )

            if not content.strip():
                continue

            # Get agenda for this decision
            agenda = self._get_agenda_for_index(candidate.keyword_index)

            # Generate decision ID
            counter_key = (candidate.decision_type, agenda)
            counters[counter_key] = counters.get(counter_key, 0) + 1
            count = counters[counter_key]

            prefix = self._get_decision_prefix(candidate.decision_type)
            meeting_num = meeting_id.replace("RAN1#", "")
            decision_id = f"{prefix}-{meeting_num}-{agenda}-{count:03d}"

            # Determine hasConsensus for Conclusion type
            # Spec: hasConsensus = true if "no consensus" is NOT in content
            has_consensus = None
            if candidate.decision_type == DecisionType.CONCLUSION:
                has_consensus = "no consensus" not in content.lower()

            # Create decision
            decision = Decision(
                decision_id=decision_id,
                decision_type=candidate.decision_type,
                meeting_id=meeting_id,
                agenda_item=agenda,
                content=content,
                paragraph_index=candidate.keyword_index,
                referenced_tdocs=TDOC_PATTERN.findall(content),
                has_ffs=bool(FFS_PATTERN.search(content)),
                has_tbd=bool(TBD_PATTERN.search(content)),
                has_consensus=has_consensus,
            )

            # Add to appropriate list
            if candidate.decision_type == DecisionType.AGREEMENT:
                agreements.append(decision)
            elif candidate.decision_type == DecisionType.CONCLUSION:
                conclusions.append(decision)
            else:
                working_assumptions.append(decision)

        return agreements, conclusions, working_assumptions

    def _find_candidates(self, paragraphs: list) -> list[DecisionCandidate]:
        """Find all decision keyword candidates.

        Supports:
        1. Header-only: "Agreement:" (content in following paragraphs)
        2. Inline: "Agreement: The UE shall..." (content on same line)

        Filters out:
        - Annex sections (per Spec requirement)
        - TOC entries (toc style paragraphs)
        """
        candidates = []
        in_annex = False

        for idx, para in enumerate(paragraphs):
            text = para.text.strip()
            style = para.style.name.lower() if para.style else ""

            # Skip TOC entries (they may contain Annex references)
            if style.startswith('toc'):
                continue

            # Check for Annex start - skip all content after this
            # Only trigger on actual headings, not TOC references
            if ANNEX_PATTERN.match(text) and style.startswith('heading'):
                in_annex = True
                continue

            # Skip if we're in Annex section
            if in_annex:
                continue

            # Try header-only patterns first
            for decision_type, pattern in DECISION_PATTERNS_HEADER.items():
                if pattern.match(text):
                    candidates.append(DecisionCandidate(
                        decision_type=decision_type,
                        keyword_index=idx,
                        keyword_text=text,
                        inline_content=None
                    ))
                    break
            else:
                # Try inline patterns (Agreement: [content])
                for decision_type, pattern in DECISION_PATTERNS_INLINE.items():
                    match = pattern.match(text)
                    if match:
                        inline_content = match.group(1).strip()
                        candidates.append(DecisionCandidate(
                            decision_type=decision_type,
                            keyword_index=idx,
                            keyword_text=text,
                            inline_content=inline_content
                        ))
                        break

        return candidates

    def _find_content_boundary(
        self,
        start_index: int,
        paragraphs: list,
        max_index: int
    ) -> int:
        """
        Find where the decision content ends.

        Ends at:
        - Next decision keyword (header or inline)
        - Heading paragraph
        - Annex section
        - Double blank lines
        """
        blank_count = 0

        for idx in range(start_index + 1, max_index):
            para = paragraphs[idx]
            text = para.text.strip()
            style = para.style.name.lower() if para.style else ""

            # Check for heading
            if style.startswith('heading'):
                return idx

            # Check for Annex (stop at Annex)
            if ANNEX_PATTERN.match(text):
                return idx

            # Check for decision keyword (header-only)
            for pattern in DECISION_PATTERNS_HEADER.values():
                if pattern.match(text):
                    return idx

            # Check for decision keyword (inline)
            for pattern in DECISION_PATTERNS_INLINE.values():
                if pattern.match(text):
                    return idx

            # Check for consecutive blanks
            if not text:
                blank_count += 1
                if blank_count >= 2:
                    return idx - 1
            else:
                blank_count = 0

        return max_index

    def _extract_content(
        self,
        paragraphs: list,
        start_index: int,
        end_index: int
    ) -> str:
        """Extract and join content paragraphs."""
        content_parts = []

        for idx in range(start_index, end_index):
            text = paragraphs[idx].text.strip()
            if text:
                content_parts.append(text)

        return '\n'.join(content_parts)

    def _get_agenda_for_index(self, para_index: int) -> str:
        """Get agenda number for a paragraph index."""
        # Find the section containing this index
        if para_index in self._section_by_index:
            return self._section_by_index[para_index].agenda_number

        # Fallback: find closest section before this index
        closest_section = None
        for section in self.sections:
            if section.start_index <= para_index:
                closest_section = section
            else:
                break

        if closest_section:
            return closest_section.agenda_number

        return "UNKNOWN"

    def _get_decision_prefix(self, decision_type: DecisionType) -> str:
        """Get prefix for decision ID."""
        prefixes = {
            DecisionType.AGREEMENT: "AGR",
            DecisionType.CONCLUSION: "CON",
            DecisionType.WORKING_ASSUMPTION: "WA",
        }
        return prefixes[decision_type]


def extract_decisions_from_docx(
    docx_path: Path,
    sections: list[Section],
    meeting_id: str
) -> tuple[list[Decision], list[Decision], list[Decision]]:
    """
    Extract decisions from a DOCX file.

    Args:
        docx_path: Path to DOCX file
        sections: List of sections with agenda numbers
        meeting_id: Meeting ID

    Returns:
        Tuple of (agreements, conclusions, working_assumptions)
    """
    doc = Document(docx_path)
    extractor = DecisionExtractor(sections)
    return extractor.extract_decisions(doc.paragraphs, meeting_id)


if __name__ == "__main__":
    import sys
    from toc_parser import parse_toc_from_docx
    from agenda_matcher import AgendaMatcher, SectionBuilder

    if len(sys.argv) < 2:
        print("Usage: python decision_extractor.py <docx_path>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    meeting_id = "RAN1#" + docx_path.stem.split('-')[-1].replace('_hybrid', '')

    # Parse TOC and build sections
    toc_entries = parse_toc_from_docx(docx_path)
    matcher = AgendaMatcher(toc_entries)
    builder = SectionBuilder(matcher)

    # Build sections from document
    doc = Document(docx_path)
    sections = []
    for idx, para in enumerate(doc.paragraphs):
        style = para.style.name.lower() if para.style else ""
        if style.startswith('heading'):
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
            section = builder.process_heading(para.text.strip(), level, idx)
            if sections:
                sections[-1].end_index = idx
            sections.append(section)

    if sections:
        sections[-1].end_index = len(doc.paragraphs)

    # Extract decisions
    agreements, conclusions, was = extract_decisions_from_docx(docx_path, sections, meeting_id)

    print(f"Extracted from {docx_path.name}:")
    print(f"  Agreements: {len(agreements)}")
    print(f"  Conclusions: {len(conclusions)}")
    print(f"  Working Assumptions: {len(was)}")

    # Show first few
    print("\nSample Agreements:")
    for agr in agreements[:3]:
        print(f"  {agr.decision_id}: {agr.content[:80]}...")
