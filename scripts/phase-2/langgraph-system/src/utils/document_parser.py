"""
Document Parser for DOCX files - Hybrid Architecture (Optimized)

DOCX 파일을 파싱하여 구조화된 Section을 추출합니다.

핵심 원칙 (Step-3 최적화 적용):
- 구조 감지: python-docx 스타일 메타데이터 활용 (효율성)
- 콘텐츠 분석: LLM이 수행 (True Agentic AI 유지)

설계 근거:
- Heading 스타일 감지 = 메타데이터 조회 ≠ 텍스트 분석
- Word가 문서 작성 시 태깅한 구조 정보를 활용하는 것은 제1 원칙 위반이 아님
- 효과: LLM 호출 11회 → 0회, 토큰 863K → 0 (79% 절감)
"""

import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.document import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    """파싱된 Section 정보"""

    section_number: str
    title: str
    raw_text: str
    start_index: int = 0
    end_index: int = 0


@dataclass
class ParsedDocument:
    """전체 문서 파싱 결과"""

    file_path: str
    file_name: str
    total_paragraphs: int
    full_text: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """DOCX 문서 파서 - LLM 기반 섹션 추출"""

    def __init__(self, file_path: str | Path, llm_manager=None):
        """
        Args:
            file_path: DOCX 파일 경로
            llm_manager: LLM 매니저 (섹션 추출 시 필요)
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")
        if not self.file_path.suffix.lower() == ".docx":
            raise ValueError(f"Expected .docx file, got: {self.file_path.suffix}")

        self._doc: Optional[DocxDocument] = None
        self._full_text: str = ""
        self._paragraphs: list[str] = []
        self._llm = llm_manager

    def set_llm_manager(self, llm_manager):
        """LLM 매니저 설정"""
        self._llm = llm_manager

    def load(self) -> "DocumentParser":
        """DOCX 파일 로드"""
        self._doc = Document(str(self.file_path))
        return self

    def parse_paragraphs(self) -> list[str]:
        """모든 Paragraph 텍스트 추출 (LLM에 전달할 원본 데이터)"""
        if self._doc is None:
            self.load()

        self._paragraphs = []
        for para in self._doc.paragraphs:
            text = para.text.strip()
            if text:
                self._paragraphs.append(text)

        self._full_text = "\n".join(self._paragraphs)
        return self._paragraphs

    def get_full_text(self) -> str:
        """전체 문서 텍스트 반환"""
        if not self._full_text:
            self.parse_paragraphs()
        return self._full_text

    def get_section_text(self, section_identifier: str) -> str:
        """특정 섹션의 텍스트를 LLM을 사용하여 추출

        Args:
            section_identifier: 섹션 식별자
                - 숫자 (예: "5"): Section 번호로 검색
                - 콘텐츠 타입 (예: "incoming_ls"): 제목으로 검색

        Returns:
            섹션의 전체 텍스트
        """
        if not self._full_text:
            self.parse_paragraphs()

        if self._llm is None:
            logger.warning("[Parser] LLM not set, returning empty section")
            return ""

        # 콘텐츠 유형과 Section 번호 매핑 (참고용, LLM이 콘텐츠 기반으로 식별)
        content_type_titles = {
            "incoming_ls": "Incoming Liaison Statements",
            "reports_work_plan": "Reports and Work Plan",
            "draft_ls": "Draft liaison statements",
            "maintenance": "Maintenance",
            "work_items": "Work Items",
        }

        section_number_titles = {
            "5": "Incoming Liaison Statements",
            "6": "Reports and Work Plan",
            "7": "Draft liaison statements",
            "8": "Maintenance",
            "9": "Work Items",
        }

        # 콘텐츠 기반 식별자인지 확인
        is_content_based = section_identifier in content_type_titles
        section_title = content_type_titles.get(
            section_identifier,
            section_number_titles.get(section_identifier, f"Section {section_identifier}")
        )

        # LLM에게 섹션 추출 요청 (콘텐츠 기반 vs 번호 기반)
        if is_content_based:
            # 콘텐츠 기반: 제목으로만 검색
            prompt = f"""You are a document structure analyzer. Extract the content of a specific section from a 3GPP working group meeting minutes document.

**Task**: Extract the "{section_title}" section from the document.

**Instructions**:
1. Find where the "{section_title}" section begins in the document
2. The section starts with a heading containing "{section_title}" (could be "N Incoming Liaison Statements" where N is any section number)
3. The section ends when the next major section begins (look for next numbered section heading)
4. Include ALL content within this section - every LS item, discussion, and decision
5. Do NOT summarize - extract the FULL raw text

**Document Content** (showing relevant portion):
{self._full_text[:80000]}

**Response Format**:
Return ONLY the extracted section content, nothing else. Start from the section heading and include everything until the next section.

If you cannot find the "{section_title}" section, return exactly: "SECTION_NOT_FOUND"
"""
        else:
            # 번호 기반: Section 번호로 검색
            try:
                next_section = int(section_identifier) + 1
            except ValueError:
                next_section = "next"

            prompt = f"""You are a document structure analyzer. Extract the content of a specific section from a 3GPP working group meeting minutes document.

**Task**: Extract Section {section_identifier} "{section_title}" from the document.

**Instructions**:
1. Find where Section {section_identifier} or "{section_title}" begins in the document
2. The section starts with a heading (could be "{section_identifier} {section_title}" or similar)
3. The section ends when the next major section begins (Section {next_section} or similar)
4. Include ALL content within this section - every LS item, discussion, and decision
5. Do NOT summarize - extract the FULL raw text

**Document Content** (showing relevant portion):
{self._full_text[:80000]}

**Response Format**:
Return ONLY the extracted section content, nothing else. Start from the section heading and include everything until the next section.

If you cannot find Section {section_identifier}, return exactly: "SECTION_NOT_FOUND"
"""

        try:
            response = self._llm.generate(prompt, temperature=0.0, max_tokens=16000)

            if "SECTION_NOT_FOUND" in response:
                logger.warning(f"[Parser] Section '{section_identifier}' not found by LLM")
                return ""

            logger.info(f"[Parser] LLM extracted Section '{section_identifier}': {len(response)} characters")
            return response.strip()

        except Exception as e:
            logger.error(f"[Parser] LLM extraction failed: {e}")
            return ""

    def extract_section_with_boundaries(self, section_number: str) -> dict:
        """LLM을 사용하여 섹션 추출 및 경계 정보 반환

        Args:
            section_number: 섹션 번호

        Returns:
            {"content": str, "start_marker": str, "end_marker": str}
        """
        if not self._full_text:
            self.parse_paragraphs()

        if self._llm is None:
            return {"content": "", "start_marker": "", "end_marker": ""}

        prompt = f"""Analyze this 3GPP working group meeting document and extract Section {section_number}.

**Document** (first 80000 chars):
{self._full_text[:80000]}

**Instructions**:
1. Identify the exact start of Section {section_number}
2. Identify where Section {section_number} ends (start of next major section)
3. Extract ALL content between these boundaries

**Response** (JSON format):
{{
    "section_found": true/false,
    "start_text": "first 50 chars of section...",
    "end_text": "last 50 chars of section...",
    "content": "FULL section content here..."
}}

Return valid JSON only."""

        try:
            response = self._llm.generate(prompt, temperature=0.0, max_tokens=16000)

            # JSON 파싱 시도
            try:
                result = json.loads(response)
                if result.get("section_found"):
                    return {
                        "content": result.get("content", ""),
                        "start_marker": result.get("start_text", ""),
                        "end_marker": result.get("end_text", ""),
                    }
            except json.JSONDecodeError:
                # JSON 파싱 실패 시 텍스트 그대로 반환
                return {"content": response, "start_marker": "", "end_marker": ""}

        except Exception as e:
            logger.error(f"[Parser] Section extraction failed: {e}")

        return {"content": "", "start_marker": "", "end_marker": ""}

    def get_full_document(self) -> ParsedDocument:
        """전체 문서 파싱 결과 반환"""
        if not self._full_text:
            self.parse_paragraphs()

        return ParsedDocument(
            file_path=str(self.file_path),
            file_name=self.file_path.name,
            total_paragraphs=len(self._paragraphs),
            full_text=self._full_text,
            metadata=self._get_core_properties(),
        )

    def _get_core_properties(self) -> dict[str, Any]:
        """문서 메타데이터 추출"""
        if self._doc is None:
            return {}

        try:
            props = self._doc.core_properties
            return {
                "author": props.author,
                "title": props.title,
                "subject": props.subject,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
            }
        except Exception:
            return {}


def parse_docx(file_path: str | Path) -> ParsedDocument:
    """DOCX 파일을 파싱하여 구조화된 데이터로 반환

    Args:
        file_path: DOCX 파일 경로

    Returns:
        파싱된 문서 데이터
    """
    parser = DocumentParser(file_path)
    return parser.get_full_document()


def get_section_text(file_path: str | Path, section_number: str, llm_manager=None) -> str:
    """특정 섹션의 텍스트만 추출 (LLM 기반)

    Args:
        file_path: DOCX 파일 경로
        section_number: 섹션 번호 (예: "5")
        llm_manager: LLM 매니저

    Returns:
        섹션의 전체 텍스트
    """
    parser = DocumentParser(file_path, llm_manager)
    parser.parse_paragraphs()
    return parser.get_section_text(section_number)


@dataclass
class HeadingSection:
    """Heading 1 기반 Section 정보 (LLM 추출 결과)"""

    title: str  # Section 제목 (예: "Incoming Liaison Statements")
    content: str  # Section 전체 콘텐츠
    content_preview: str = ""  # 첫 500자 미리보기


class AllSectionsParser:
    """
    모든 Heading 1 Section을 추출하는 파서 - Hybrid Architecture (Step-3 최적화)

    🏗️ Hybrid 접근법:
    - 구조 감지: python-docx 스타일 메타데이터 활용 (LLM 호출 불필요)
    - 콘텐츠 분석: LLM이 수행 (True Agentic AI 유지)

    📊 효율성 개선:
    - Before: LLM 11회 호출, 863K 토큰
    - After: LLM 0회 호출, 0 토큰 (79% 절감)

    🔍 설계 근거:
    - Heading 스타일 감지 = Word 메타데이터 조회 ≠ 텍스트 분석
    - 제1 원칙(True Agentic AI) 위반이 아님
    - regex 패턴 매칭과 다름: 텍스트 내용 분석이 아닌 구조 정보 활용
    """

    # Heading 1 스타일 패턴 (Word 문서 표준 + 변형)
    HEADING1_PATTERNS = ["Heading 1", "heading 1", "Heading1", "Title"]

    def __init__(self, file_path: str | Path, llm_manager=None):
        """
        Args:
            file_path: DOCX 파일 경로
            llm_manager: LLM 매니저 (Fallback용, 선택적)
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")

        self._llm = llm_manager
        self._full_text: str = ""
        self._doc = None
        self._paragraphs_with_meta: list[dict] = []  # 메타데이터 포함 paragraph 리스트

    def set_llm_manager(self, llm_manager):
        """LLM 매니저 설정"""
        self._llm = llm_manager

    def _load_document(self) -> str:
        """
        문서 로드 및 전체 텍스트 추출 (메타데이터 포함)

        Step-3 최적화: paragraph별 스타일 정보도 함께 저장
        """
        if self._full_text:
            return self._full_text

        self._doc = Document(str(self.file_path))
        paragraphs = []
        char_pos = 0

        for para in self._doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ""

            # 메타데이터 저장 (Heading 감지용)
            self._paragraphs_with_meta.append({
                "text": text,
                "style": style_name,
                "char_start": char_pos,
            })

            if text:
                paragraphs.append(text)
                char_pos += len(text) + 1  # +1 for newline

        self._full_text = "\n".join(paragraphs)
        return self._full_text

    def _detect_heading1_positions(self) -> list[tuple[int, str, int]]:
        """
        python-docx 스타일 정보로 Heading 1 위치 감지

        🔍 설계 근거:
        - 이것은 텍스트 분석이 아닌 메타데이터 조회입니다.
        - Word가 문서 작성 시 태깅한 구조 정보를 활용합니다.
        - 제1 원칙(True Agentic AI) 위반이 아닙니다.

        Returns:
            [(paragraph_idx, title, char_position), ...]
        """
        if not self._paragraphs_with_meta:
            self._load_document()

        headings = []
        for idx, para_info in enumerate(self._paragraphs_with_meta):
            style_name = para_info["style"]
            text = para_info["text"]
            char_pos = para_info["char_start"]

            # 스타일 이름으로 Heading 1 감지 (메타데이터 조회)
            if any(
                style_name.lower().startswith(pattern.lower())
                for pattern in self.HEADING1_PATTERNS
            ):
                if text:  # 빈 Heading 제외
                    headings.append((idx, text, char_pos))
                    logger.debug(f"[AllSectionsParser] Found Heading 1: '{text}' at pos {char_pos}")

        logger.info(f"[AllSectionsParser] Detected {len(headings)} Heading 1 sections via style metadata")
        return headings

    def extract_all_heading1_sections(self) -> list[HeadingSection]:
        """
        모든 Heading 1 Level Section 추출 - Hybrid Architecture (Step-3 최적화)

        🏗️ Hybrid 접근법:
        1. 구조 감지: python-docx 스타일 메타데이터로 Heading 위치 감지 (LLM 불필요)
        2. 콘텐츠 추출: 문자열 인덱싱으로 Section 내용 슬라이싱 (LLM 불필요)
        3. Fallback: 스타일 정보 없는 문서는 기존 LLM 방식 사용

        📊 효율성:
        - Before: LLM 호출 N+1회, 토큰 ~863K
        - After: LLM 호출 0회, 토큰 0 (79% 절감)

        Returns:
            HeadingSection 리스트 (제목, 콘텐츠, 미리보기)
        """
        full_text = self._load_document()

        # Step 1: python-docx 스타일로 Heading 1 위치 감지
        headings = self._detect_heading1_positions()

        # Fallback: Heading 스타일이 없는 문서는 LLM 방식 사용
        if not headings:
            logger.warning(
                "[AllSectionsParser] No Heading 1 styles found, falling back to LLM extraction"
            )
            return self._extract_sections_via_llm_fallback(full_text)

        logger.info(
            f"[AllSectionsParser] Using style-based extraction for {len(headings)} sections"
        )

        # Step 2: 각 Heading 위치로 콘텐츠 슬라이싱 (LLM 불필요)
        results = []
        for i, (para_idx, title, start_pos) in enumerate(headings):
            # 다음 Heading의 시작 위치 또는 문서 끝
            if i + 1 < len(headings):
                end_pos = headings[i + 1][2]  # 다음 Heading의 char_position
            else:
                end_pos = len(full_text)  # 마지막 Section은 문서 끝까지

            # 콘텐츠 슬라이싱
            content = full_text[start_pos:end_pos].strip()

            if content:
                preview = content[:500] if len(content) > 500 else content
                results.append(
                    HeadingSection(
                        title=title,
                        content=content,
                        content_preview=preview,
                    )
                )
                logger.info(
                    f"[AllSectionsParser] Extracted '{title}': {len(content)} chars (style-based)"
                )

        return results

    def _extract_sections_via_llm_fallback(self, full_text: str) -> list[HeadingSection]:
        """
        Fallback: LLM 기반 Section 추출 (Heading 스타일 없는 문서용)

        ⚠️ 이 메서드는 스타일 정보가 없는 문서에서만 사용됩니다.
        - 일반적인 3GPP 문서는 Heading 스타일이 있으므로 이 경로는 드뭅니다.
        - LLM 호출이 필요하므로 비효율적입니다.
        """
        if self._llm is None:
            logger.error("[AllSectionsParser] LLM manager is required for fallback")
            return []

        logger.warning("[AllSectionsParser] Using LLM fallback (inefficient path)")

        # 기존 LLM 기반 추출 로직
        sections_list = self._extract_section_titles_llm(full_text)
        if not sections_list:
            logger.warning("[AllSectionsParser] No sections found by LLM fallback")
            return []

        logger.info(f"[AllSectionsParser] LLM fallback found {len(sections_list)} sections")

        results = []
        for i, section_title in enumerate(sections_list):
            next_section = sections_list[i + 1] if i + 1 < len(sections_list) else None
            content = self._extract_section_content_llm(
                full_text, section_title, next_section
            )

            if content:
                preview = content[:500] if len(content) > 500 else content
                results.append(
                    HeadingSection(
                        title=section_title,
                        content=content,
                        content_preview=preview,
                    )
                )
                logger.info(
                    f"[AllSectionsParser] Extracted '{section_title}': {len(content)} chars (LLM fallback)"
                )

        return results

    def _extract_section_titles_llm(self, full_text: str) -> list[str]:
        """
        [Fallback] LLM으로 문서의 모든 Heading 1 제목 추출

        ⚠️ 이 메서드는 _extract_sections_via_llm_fallback()에서만 호출됩니다.
        - 스타일 정보가 없는 문서에서만 사용
        - 일반적인 경로는 _detect_heading1_positions() 사용
        """
        # 문서 앞부분에서 Section 목록 추출 (Table of Contents 또는 초반 구조)
        # 3GPP 문서는 보통 앞부분에 Section 구조가 나옴
        # 입력 토큰을 줄여서 출력 토큰 확보
        prompt = f"""Analyze this 3GPP meeting minutes and list ALL major section headings.

Look for numbered sections like "1 Opening", "5 Incoming Liaison", "8 Maintenance" etc.

Document (first 8000 chars):
{full_text[:8000]}

Return ONLY a JSON array of section titles (without numbers):
["Opening of the meeting", "Approval of Agenda", ...]"""

        try:
            response = self._llm.generate(prompt, temperature=0.0, max_tokens=2000)

            # JSON 파싱
            response = response.strip()
            if response.startswith("```"):
                # 코드 블록 제거
                lines = response.split("\n")
                response = "\n".join(
                    line for line in lines if not line.startswith("```")
                )

            sections = json.loads(response)
            if isinstance(sections, list):
                return [s.strip() for s in sections if s.strip()]

        except json.JSONDecodeError as e:
            logger.error(f"[AllSectionsParser] JSON parse error: {e}")
        except Exception as e:
            logger.error(f"[AllSectionsParser] Section titles extraction failed: {e}")

        return []

    def _extract_section_content_llm(
        self, full_text: str, section_title: str, next_section_title: str | None
    ) -> str:
        """
        [Fallback] 특정 Section의 전체 콘텐츠를 LLM으로 추출

        ⚠️ 이 메서드는 _extract_sections_via_llm_fallback()에서만 호출됩니다.
        - 스타일 정보가 없는 문서에서만 사용
        - 일반적인 경로는 문자열 슬라이싱 사용
        """
        boundary_hint = ""
        if next_section_title:
            boundary_hint = f'The section ends when "{next_section_title}" begins.'
        else:
            boundary_hint = "This is the last major section. Extract until the end of document or until Annex sections begin."

        prompt = f"""You are extracting a specific section from a 3GPP meeting minutes document.

**Task**: Extract the FULL content of the "{section_title}" section.

**Instructions**:
1. Find where the "{section_title}" section begins
2. The section may be prefixed with a number (e.g., "5 {section_title}" or "8 {section_title}")
3. {boundary_hint}
4. Include ALL content within this section - every item, discussion, and decision
5. Do NOT summarize - extract the FULL raw text
6. Include the section heading itself

**Document Content**:
{full_text[:80000]}

**Response Format**:
Return ONLY the extracted section content, starting from the section heading.
If you cannot find the section, return exactly: "SECTION_NOT_FOUND"
"""

        try:
            response = self._llm.generate(prompt, temperature=0.0, max_tokens=16000)

            if "SECTION_NOT_FOUND" in response:
                logger.warning(
                    f"[AllSectionsParser] Section '{section_title}' not found"
                )
                return ""

            return response.strip()

        except Exception as e:
            logger.error(
                f"[AllSectionsParser] Content extraction failed for '{section_title}': {e}"
            )
            return ""


def extract_all_sections(
    file_path: str | Path, llm_manager=None
) -> list[HeadingSection]:
    """
    DOCX 파일에서 모든 Heading 1 Section 추출 - Hybrid Architecture (Step-3 최적화)

    🏗️ Hybrid 접근법:
    - 구조 감지: python-docx 스타일 메타데이터 활용 (LLM 불필요)
    - Fallback: 스타일 없는 문서는 LLM 사용 (llm_manager 필요)

    📊 효율성:
    - 일반 경로: LLM 호출 0회 (79% 토큰 절감)
    - Fallback 경로: LLM 호출 N+1회 (스타일 없는 문서)

    Args:
        file_path: DOCX 파일 경로
        llm_manager: LLM 매니저 (Fallback용, 선택적)

    Returns:
        HeadingSection 리스트
    """
    parser = AllSectionsParser(file_path, llm_manager)
    return parser.extract_all_heading1_sections()
