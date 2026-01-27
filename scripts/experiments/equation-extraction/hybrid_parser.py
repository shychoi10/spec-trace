"""
하이브리드 파서 V2: Phase-3 python-docx 텍스트 + V2.1 OMML 수식 추출

핵심 수정 (V2):
- content_raw: para._element에서 w:r + m:oMath 텍스트 모두 포함
- offset 계산을 content_raw 기준으로 수행 (para.text 기반 content에는 수식 텍스트 없음)
- Phase-3 호환: content (para.text 기반) 유지 + content_raw (수식 포함) 추가

순서:
- Phase-3 ReportParser로 decision 추출 (텍스트 동일 보장)
- python-docx paragraph XML에서 OMML 수식 추출
- content_raw 구축 (w:r + m:oMath 텍스트 통합)
- Decision content에 수식 매핑 + content_raw 기준 offset 재계산
- 기존 Phase-3 결과와 비교
"""

import sys
import json
import re
import zipfile
import tempfile
import subprocess
import os
import shutil
import time
from pathlib import Path
from datetime import datetime
from typing import Tuple, List, Optional
from concurrent.futures import ThreadPoolExecutor

from docx import Document
from lxml import etree

# Phase-3 parser imports
PHASE3_DIR = Path(__file__).resolve().parent.parent.parent / 'phase-3' / 'report-parser'
sys.path.insert(0, str(PHASE3_DIR))
from models import DecisionType, ParsedReport
from report_parser import ReportParser

# Namespace URIs
MATH_URI = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_URI = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Pandoc temp DOCX templates
CONTENT_TYPES_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''

RELS_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''

DOCUMENT_TEMPLATE = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
    <w:body>
        <w:p>
            {content}
        </w:p>
    </w:body>
</w:document>'''

# Inline content patterns (same as Phase-3 decision_extractor)
INLINE_PATTERNS = [
    re.compile(r'^agreements?\s*:\s*(.+)', re.IGNORECASE | re.DOTALL),
    re.compile(r'^conclusions?\s*:\s*(.+)', re.IGNORECASE | re.DOTALL),
    re.compile(r'^working\s*assumptions?\s*:\s*(.+)', re.IGNORECASE | re.DOTALL),
]


# =============================================================================
# V2.1 Batch Equation Converter (from full_parse_test_v2.py)
# =============================================================================

class BatchEquationConverter:
    """배치 수식 변환기"""

    def __init__(self, max_workers: int = 4):
        self.max_workers = max_workers
        self._temp_base = None

    def cleanup(self):
        if self._temp_base and self._temp_base.exists():
            shutil.rmtree(self._temp_base, ignore_errors=True)
        self._temp_base = None

    def _setup_temp_dir(self):
        if self._temp_base is None:
            self._temp_base = Path(tempfile.mkdtemp())
        return self._temp_base

    def _create_temp_docx(self, omml_xml: str, index: int) -> Path:
        temp_dir = self._setup_temp_dir()
        temp_docx = temp_dir / f"eq_{index}.docx"
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('[Content_Types].xml', CONTENT_TYPES_XML)
            zf.writestr('_rels/.rels', RELS_XML)
            zf.writestr('word/document.xml', DOCUMENT_TEMPLATE.format(content=omml_xml))
        return temp_docx

    def _convert_single(self, args: Tuple[int, str]) -> Tuple[int, str, bool, str]:
        index, omml_xml = args
        temp_docx = self._create_temp_docx(omml_xml, index)
        try:
            result = subprocess.run(
                ['pandoc', str(temp_docx), '-t', 'latex', '--wrap=none'],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                latex = result.stdout.strip()
                is_valid, note = self._validate_latex(latex)
                return (index, latex, is_valid, note)
            else:
                return (index, f"[ERROR: {result.stderr[:50]}]", False, "Pandoc error")
        except Exception as e:
            return (index, f"[ERROR: {str(e)[:50]}]", False, str(e))
        finally:
            try:
                os.unlink(temp_docx)
            except:
                pass

    def _validate_latex(self, latex: str) -> Tuple[bool, str]:
        import re as _re
        # Step 1: \left/\right delimiter 중립화
        # \left( 등의 괄호는 구분자(delimiter)이므로 그루핑 괄호와 분리하여 카운트
        cleaned = latex
        cleaned = _re.sub(r'\\left\s*[\(\[\{|.]', ' LDEL ', cleaned)
        cleaned = _re.sub(r'\\left\\[a-zA-Z]+', ' LDEL ', cleaned)
        cleaned = _re.sub(r'\\right\s*[\)\]\}|.]', ' RDEL ', cleaned)
        cleaned = _re.sub(r'\\right\\[a-zA-Z]+', ' RDEL ', cleaned)

        # Step 2: 괄호 유형별 불균형 계산
        p_diff = abs(cleaned.count('(') - cleaned.count(')'))
        b_diff = abs(cleaned.count('{') - cleaned.count('}'))
        k_diff = abs(cleaned.count('[') - cleaned.count(']'))

        # Step 3: OMML boundary artifact 허용
        # 3GPP DOCX에서 수식의 괄호가 일반 텍스트(w:r)와 OMML(m:oMath) 사이에
        # 분리되어 Pandoc 변환 시 괄호 짝이 안 맞을 수 있음
        # - paren: OMML 경계 영향 (≤5 허용)
        # - brace: LaTeX 구문 필수이므로 엄격 (≤1 허용)
        # - bracket: OMML 경계 영향 (≤1 허용)
        if p_diff > 5 or b_diff > 1 or k_diff > 1:
            return False, f"Unbalanced brackets: paren±{p_diff} brace±{b_diff} bracket±{k_diff}"

        if '[ERROR' in latex or latex.startswith('['):
            return False, "Contains error marker"

        if p_diff + b_diff + k_diff > 0:
            return True, f"OMML boundary artifact: paren±{p_diff} brace±{b_diff} bracket±{k_diff}"
        return True, ""

    def convert_batch(self, equations: List[Tuple[int, str]]) -> dict:
        if not equations:
            return {}
        results = {}
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = list(executor.map(self._convert_single, equations))
        for index, latex, is_valid, note in futures:
            results[index] = {'latex': latex, 'is_valid': is_valid, 'validation_note': note}
        return results


# =============================================================================
# OMML Extraction Helpers (lxml-based, for python-docx para._element)
# =============================================================================

def extract_omml_text(omml_element) -> str:
    """lxml OMML 요소에서 텍스트 추출"""
    texts = []
    for elem in omml_element.iter():
        if elem.text:
            texts.append(elem.text)
    return ''.join(texts)


def omml_to_xml_string(omml_element) -> str:
    """lxml element → XML string"""
    return etree.tostring(omml_element, encoding='unicode')


def is_block_equation(omath_element) -> bool:
    """블록 수식 여부 (부모가 oMathPara인지)"""
    parent = omath_element.getparent()
    if parent is not None:
        tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
        return tag == 'oMathPara'
    return False


# =============================================================================
# Full Text Extraction (w:r + m:oMath + m:oMathPara + w:hyperlink)
# =============================================================================

def _collect_run_text(run_element, parts: list, w_ns: str):
    """Collect text from a w:r element (handles w:t, w:tab, w:br)."""
    for child in run_element:
        tag = child.tag
        if tag == f'{w_ns}t':
            if child.text:
                parts.append(child.text)
        elif tag == f'{w_ns}tab':
            parts.append('\t')
        elif tag == f'{w_ns}br':
            parts.append('\n')


def extract_paragraph_full_text(para_element) -> str:
    """
    Extract full text from paragraph element including equation text.

    Unlike python-docx's para.text which only collects w:r/w:t text,
    this also includes m:oMath and m:oMathPara text.

    Handles:
    - w:r → w:t (text), w:tab (→ tab), w:br (→ newline)
    - w:hyperlink → w:r → w:t
    - m:oMath (inline equation text)
    - m:oMathPara → m:oMath (block equation text)
    - w:ins (tracked changes insertions) → w:r → w:t
    """
    parts = []
    w_ns = f'{{{WORD_URI}}}'
    m_ns = f'{{{MATH_URI}}}'

    for child in para_element:
        tag = child.tag

        if tag == f'{w_ns}r':
            # Normal text run
            _collect_run_text(child, parts, w_ns)

        elif tag == f'{w_ns}hyperlink':
            # Hyperlink: contains w:r children
            for r in child.findall(f'{w_ns}r'):
                _collect_run_text(r, parts, w_ns)

        elif tag == f'{m_ns}oMath':
            # Inline equation
            parts.append(extract_omml_text(child))

        elif tag == f'{m_ns}oMathPara':
            # Block equation container
            for omath in child.findall(f'{m_ns}oMath'):
                parts.append(extract_omml_text(omath))

        elif tag == f'{w_ns}ins':
            # Tracked changes (insertions)
            for r in child.findall(f'{w_ns}r'):
                _collect_run_text(r, parts, w_ns)

    return ''.join(parts)


# =============================================================================
# Hybrid Parser
# =============================================================================

class HybridReportParser:
    """
    하이브리드 파서 V2: Phase-3 텍스트 추출 + V2.1 수식 추출

    텍스트: python-docx para.text (Phase-3 동일) → content
    수식 포함 텍스트: para._element full extraction → content_raw
    수식: para._element에서 m:oMath → Pandoc batch → LaTeX
    offset: content_raw 기준 계산 (para.text에는 수식 텍스트 없음)
    """

    def __init__(self, docx_path: Path, meeting_id: Optional[str] = None):
        self.docx_path = docx_path
        self.meeting_id = meeting_id
        self.converter = BatchEquationConverter()

    def parse(self):
        """
        Returns:
            report: Phase-3 ParsedReport (텍스트 동일)
            enriched: 수식 포함 Decision 리스트 (content + content_raw)
            eq_stats: 수식 통계
        """
        start = time.time()

        # Step 1: Phase-3 파서로 decision 추출
        print("  [1/5] Phase-3 parser: decision extraction...")
        phase3_parser = ReportParser(self.docx_path, self.meeting_id)
        report = phase3_parser.parse()
        print(f"        → {report.total_decisions} decisions "
              f"({len(report.agreements)} AGR, {len(report.conclusions)} CON, "
              f"{len(report.working_assumptions)} WA)")

        # Step 2: DOCX paragraph별 수식 추출
        print("  [2/5] Extracting OMML equations from paragraph XML...")
        doc = Document(self.docx_path)
        para_equations, pending = self._extract_paragraph_equations(doc)
        total_eq = sum(len(eqs) for eqs in para_equations.values())
        print(f"        → {total_eq} equations in {len(para_equations)} paragraphs")

        # Step 3: Batch LaTeX 변환
        print(f"  [3/5] Batch LaTeX conversion ({len(pending)} equations)...")
        conversion_results = self.converter.convert_batch(pending)
        self.converter.cleanup()

        # Step 4: Decision에 수식 매핑 + content_raw 구축
        print("  [4/5] Mapping equations to decisions + building content_raw...")
        all_decisions = report.agreements + report.conclusions + report.working_assumptions
        enriched = self._map_equations_to_decisions(
            doc, all_decisions, para_equations, conversion_results
        )

        # Step 5: 통계
        eq_stats = self._compute_stats(enriched, total_eq)
        elapsed = time.time() - start
        eq_stats['elapsed'] = round(elapsed, 1)

        print(f"  [5/5] Done in {elapsed:.1f}s")
        print(f"        → {eq_stats['total_equations_in_decisions']} equations in decisions, "
              f"{eq_stats['valid']} valid, {eq_stats['invalid']} invalid")
        print(f"        → {eq_stats['offset_exact']} offset exact, "
              f"{eq_stats['offset_failed']} offset failed")

        return report, enriched, eq_stats

    def _extract_paragraph_equations(self, doc):
        """각 paragraph의 XML에서 OMML 수식 추출"""
        para_equations = {}  # para_index → [eq_info, ...]
        pending = []  # (global_idx, omml_xml) for batch conversion

        ns = f'{{{MATH_URI}}}'

        for idx, para in enumerate(doc.paragraphs):
            elem = para._element  # lxml element
            omaths = elem.findall(f'.//{ns}oMath')

            if not omaths:
                continue

            eqs = []
            for om in omaths:
                plain = extract_omml_text(om)
                xml_str = omml_to_xml_string(om)
                display = 'block' if is_block_equation(om) else 'inline'

                global_idx = len(pending)
                pending.append((global_idx, xml_str))

                eqs.append({
                    'plain_text': plain.strip(),
                    'display_type': display,
                    'global_idx': global_idx,
                })

            para_equations[idx] = eqs

        return para_equations, pending

    def _build_content_raw_and_range(self, paragraphs, keyword_idx, content):
        """
        Build content_raw (includes equation text) and find paragraph range.

        content_raw differs from content (Phase-3) in that it includes OMML
        equation text, which para.text excludes.

        Returns:
            (content_raw, para_indices)
        """
        content_lines = [l for l in content.split('\n') if l.strip()]
        if not content_lines:
            return content, []

        # Check for inline content
        keyword_text = paragraphs[keyword_idx].text.strip()
        is_inline = False
        inline_full_text = None

        for pat in INLINE_PATTERNS:
            m = pat.match(keyword_text)
            if m:
                is_inline = True
                # Extract inline portion from full_text (includes equation text)
                keyword_ft = extract_paragraph_full_text(paragraphs[keyword_idx]._element)
                fm = pat.match(keyword_ft)
                if fm:
                    inline_full_text = fm.group(1).strip()
                else:
                    # Fallback: use para.text-based inline text
                    inline_full_text = m.group(1).strip()
                break

        # Find matched paragraph indices by matching para.text against content lines
        search_start = keyword_idx + 1
        matched = []
        line_idx = 0

        if is_inline:
            matched.append(keyword_idx)
            line_idx = 1  # first content line is inline text

        for idx in range(search_start, min(keyword_idx + 500, len(paragraphs))):
            if line_idx >= len(content_lines):
                break

            para_text = paragraphs[idx].text.strip()
            if not para_text:
                continue

            if para_text == content_lines[line_idx]:
                matched.append(idx)
                line_idx += 1

        if not matched:
            return content, []

        # Build content_raw from contiguous range [first_matched, last_matched]
        # This includes equation-only paragraphs between matched paragraphs
        range_start = matched[0]
        range_end = matched[-1] + 1

        raw_parts = []
        all_indices = []

        for idx in range(range_start, range_end):
            if idx == keyword_idx and is_inline:
                # Inline keyword paragraph: use inline portion of full_text
                if inline_full_text:
                    raw_parts.append(inline_full_text)
                    all_indices.append(idx)
            else:
                ft = extract_paragraph_full_text(paragraphs[idx]._element).strip()
                if ft:
                    raw_parts.append(ft)
                    all_indices.append(idx)

        content_raw = '\n'.join(raw_parts)
        return content_raw, all_indices

    def _map_equations_to_decisions(self, doc, decisions, para_equations, conversion_results):
        """Decision에 수식 매핑 + content_raw 구축 + offset 재계산"""
        enriched = []
        paragraphs = doc.paragraphs

        for d in decisions:
            # Build content_raw and get paragraph range
            content_raw, content_paras = self._build_content_raw_and_range(
                paragraphs, d.paragraph_index, d.content
            )

            # Collect equations from content paragraphs
            decision_eqs = []
            for pi in content_paras:
                if pi in para_equations:
                    for eq in para_equations[pi]:
                        r = conversion_results.get(eq['global_idx'], {
                            'latex': '[NO_CONVERSION]',
                            'is_valid': False,
                            'validation_note': 'Missing'
                        })
                        decision_eqs.append({
                            'plain_text': eq['plain_text'],
                            'latex': r['latex'],
                            'display_type': eq['display_type'],
                            'is_valid': r['is_valid'],
                            'validation_note': r['validation_note'],
                            'position': None,
                            'char_offset': None,
                            'char_length': None,
                        })

            # Offset recalculation using content_raw (not content!)
            # content (para.text) doesn't include equation text → offset fails
            # content_raw (full extraction) includes equation text → offset works
            if decision_eqs:
                self._recalculate_offsets(decision_eqs, content_raw)

            enriched.append({
                'decision_id': d.decision_id,
                'decision_type': d.decision_type.value,
                'meeting_id': d.meeting_id,
                'agenda_item': d.agenda_item,
                'content': d.content,           # Phase-3 compatible (no equation text)
                'content_raw': content_raw,      # Includes equation text (for offset)
                'paragraph_index': d.paragraph_index,
                'referenced_tdocs': d.referenced_tdocs,
                'has_ffs': d.has_ffs,
                'has_tbd': d.has_tbd,
                'equations': decision_eqs,
                'content_para_indices': content_paras,
            })

        return enriched

    def _recalculate_offsets(self, equations, content_raw):
        """Decision content_raw 내에서 수식 position/char_offset 재계산"""
        search_start = 0
        for i, eq in enumerate(equations):
            eq['position'] = i
            plain = eq['plain_text'].strip()
            eq['plain_text'] = plain
            eq['char_length'] = len(plain)

            found_at = content_raw.find(plain, search_start)
            if found_at >= 0:
                eq['char_offset'] = found_at
                search_start = found_at + len(plain)
            else:
                eq['char_offset'] = -1  # 찾지 못함

    def _compute_stats(self, enriched, total_eq_in_doc):
        """수식 통계 계산 (content_raw 기준 offset 검증)"""
        stats = {
            'total_equations_in_doc': total_eq_in_doc,
            'total_equations_in_decisions': 0,
            'valid': 0,
            'invalid': 0,
            'inline': 0,
            'block': 0,
            'offset_exact': 0,
            'offset_failed': 0,
            'decisions_with_equations': 0,
            'position_ok': 0,
            'position_fail': 0,
            'monotonic_ok': 0,
            'monotonic_fail': 0,
        }

        for d in enriched:
            eqs = d['equations']
            content_raw = d['content_raw']  # Verify against content_raw

            if eqs:
                stats['decisions_with_equations'] += 1

            for i, eq in enumerate(eqs):
                stats['total_equations_in_decisions'] += 1

                if eq['is_valid']:
                    stats['valid'] += 1
                else:
                    stats['invalid'] += 1

                if eq['display_type'] == 'inline':
                    stats['inline'] += 1
                else:
                    stats['block'] += 1

                # offset 검증 against content_raw
                if eq['char_offset'] >= 0:
                    segment = content_raw[eq['char_offset']:eq['char_offset'] + eq['char_length']]
                    if segment == eq['plain_text']:
                        stats['offset_exact'] += 1
                    else:
                        stats['offset_failed'] += 1
                else:
                    stats['offset_failed'] += 1

                # position 검증
                if eq['position'] == i:
                    stats['position_ok'] += 1
                else:
                    stats['position_fail'] += 1

            # monotonic 검증
            for i in range(1, len(eqs)):
                if eqs[i]['char_offset'] >= 0 and eqs[i - 1]['char_offset'] >= 0:
                    if eqs[i]['char_offset'] > eqs[i - 1]['char_offset']:
                        stats['monotonic_ok'] += 1
                    else:
                        stats['monotonic_fail'] += 1

        return stats


# =============================================================================
# Phase-3 비교
# =============================================================================

def compare_with_phase3(enriched, phase3_path: Path):
    """Phase-3 기존 결과와 엄밀 비교"""
    with open(phase3_path, 'r', encoding='utf-8') as f:
        phase3 = json.load(f)

    # Phase-3 decisions를 ID → content 맵으로
    phase3_decisions = {}
    for dtype in ['agreements', 'conclusions', 'working_assumptions']:
        for d in phase3.get(dtype, []):
            phase3_decisions[d['decision_id']] = d['content']

    # Hybrid decisions를 ID → content 맵으로
    hybrid_decisions = {}
    for d in enriched:
        hybrid_decisions[d['decision_id']] = d['content']

    # 비교
    results = {
        'phase3_total': len(phase3_decisions),
        'hybrid_total': len(hybrid_decisions),
        'common': 0,
        'content_exact_match': 0,
        'content_mismatch': 0,
        'phase3_only': [],
        'hybrid_only': [],
        'mismatches': [],  # 처음 10개만
    }

    common_ids = set(phase3_decisions.keys()) & set(hybrid_decisions.keys())
    results['common'] = len(common_ids)
    results['phase3_only'] = sorted(set(phase3_decisions.keys()) - set(hybrid_decisions.keys()))
    results['hybrid_only'] = sorted(set(hybrid_decisions.keys()) - set(phase3_decisions.keys()))

    for did in sorted(common_ids):
        p3_content = phase3_decisions[did]
        hy_content = hybrid_decisions[did]
        if p3_content == hy_content:
            results['content_exact_match'] += 1
        else:
            results['content_mismatch'] += 1
            if len(results['mismatches']) < 10:
                diff_pos = next(
                    (i for i, (a, b) in enumerate(zip(p3_content, hy_content)) if a != b),
                    min(len(p3_content), len(hy_content))
                )
                results['mismatches'].append({
                    'decision_id': did,
                    'diff_position': diff_pos,
                    'phase3_len': len(p3_content),
                    'hybrid_len': len(hy_content),
                    'phase3_snippet': p3_content[max(0, diff_pos - 20):diff_pos + 40],
                    'hybrid_snippet': hy_content[max(0, diff_pos - 20):diff_pos + 40],
                })

    return results


# =============================================================================
# Main
# =============================================================================

def main():
    docx_path = Path("/home/sihyeon/workspace/spec-trace/ontology/input/meetings/RAN1/Final_Report/Final_Report_RAN1-112.docx")
    phase3_path = Path("/home/sihyeon/workspace/spec-trace/ontology/output/parsed_reports/RAN1_112.json")
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)

    print("=" * 80)
    print("하이브리드 파서 V2 검증: Phase-3 텍스트 + V2.1 수식 + content_raw offset")
    print("=" * 80)
    print(f"입력: {docx_path.name}")
    print()

    # 1. 하이브리드 파싱
    parser = HybridReportParser(docx_path)
    report, enriched, eq_stats = parser.parse()

    # 2. Phase-3 비교
    print()
    print("=" * 80)
    print("Phase-3 기존 결과와 텍스트 비교 (content = para.text 기반)")
    print("=" * 80)

    comparison = compare_with_phase3(enriched, phase3_path)
    print(f"  Phase-3 decisions: {comparison['phase3_total']}")
    print(f"  Hybrid decisions:  {comparison['hybrid_total']}")
    print(f"  공통 ID:           {comparison['common']}")
    print(f"  텍스트 완전 일치:  {comparison['content_exact_match']}/{comparison['common']}")
    print(f"  텍스트 불일치:     {comparison['content_mismatch']}")
    print(f"  Phase-3에만:       {len(comparison['phase3_only'])}")
    print(f"  Hybrid에만:        {len(comparison['hybrid_only'])}")

    if comparison['mismatches']:
        print()
        print("  --- 텍스트 불일치 상세 (처음 10개) ---")
        for m in comparison['mismatches']:
            print(f"  {m['decision_id']}: diff@{m['diff_position']} "
                  f"(p3:{m['phase3_len']} vs hy:{m['hybrid_len']})")
            print(f"    Phase-3: ...{m['phase3_snippet']!r}...")
            print(f"    Hybrid:  ...{m['hybrid_snippet']!r}...")

    # 3. 수식 통계
    print()
    print("=" * 80)
    print("수식 추출 통계 (offset = content_raw 기준)")
    print("=" * 80)
    print(f"  문서 전체 수식:     {eq_stats['total_equations_in_doc']}")
    print(f"  Decision 내 수식:   {eq_stats['total_equations_in_decisions']}")
    print(f"  Valid:              {eq_stats['valid']}")
    print(f"  Invalid:            {eq_stats['invalid']}")
    print(f"  Inline:             {eq_stats['inline']}")
    print(f"  Block:              {eq_stats['block']}")
    print(f"  Offset 정합:        {eq_stats['offset_exact']}/{eq_stats['total_equations_in_decisions']}")
    print(f"  Offset 실패:        {eq_stats['offset_failed']}")
    print(f"  Position OK:        {eq_stats['position_ok']}")
    print(f"  Monotonic OK:       {eq_stats['monotonic_ok']}")
    print(f"  수식 포함 Decision: {eq_stats['decisions_with_equations']}")
    print(f"  소요 시간:          {eq_stats['elapsed']}s")

    # 4. content vs content_raw 차이 분석
    print()
    print("=" * 80)
    print("content vs content_raw 비교 (수식 포함 확인)")
    print("=" * 80)
    content_diff_count = 0
    content_same_count = 0
    raw_longer_total = 0
    for d in enriched:
        if d['content'] != d['content_raw']:
            content_diff_count += 1
            raw_longer_total += len(d['content_raw']) - len(d['content'])
        else:
            content_same_count += 1
    print(f"  content == content_raw: {content_same_count} decisions (수식 없음)")
    print(f"  content != content_raw: {content_diff_count} decisions (수식 포함)")
    print(f"  content_raw 추가 문자:  {raw_longer_total} chars (수식 텍스트)")

    # 5. Completeness
    eq_outside = eq_stats['total_equations_in_doc'] - eq_stats['total_equations_in_decisions']
    print()
    print("=" * 80)
    print("수식 Completeness")
    print("=" * 80)
    print(f"  문서 전체:   {eq_stats['total_equations_in_doc']}")
    print(f"  Decision 내: {eq_stats['total_equations_in_decisions']}")
    print(f"  Decision 외: {eq_outside} (non-decision 단락의 수식, 정상)")

    # 6. 결과 저장
    output = {
        'meeting_id': report.meeting_id,
        'parsed_at': datetime.now().isoformat(),
        'version': 'hybrid_v2',
        'phase3_comparison': comparison,
        'equation_stats': eq_stats,
        'decisions': [
            {
                'decision_id': d['decision_id'],
                'decision_type': d['decision_type'],
                'agenda_item': d['agenda_item'],
                'content': d['content'],
                'content_raw': d['content_raw'],
                'equations': d['equations'],
            }
            for d in enriched
        ],
    }

    output_file = output_dir / "hybrid_v2_RAN1_112.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\n결과 저장: {output_file}")

    # 7. 결론
    print()
    print("=" * 80)
    print("검증 결론")
    print("=" * 80)
    text_match_pct = (comparison['content_exact_match'] / comparison['common'] * 100
                      if comparison['common'] > 0 else 0)
    eq_valid_pct = (eq_stats['valid'] / eq_stats['total_equations_in_decisions'] * 100
                    if eq_stats['total_equations_in_decisions'] > 0 else 0)
    offset_pct = (eq_stats['offset_exact'] / eq_stats['total_equations_in_decisions'] * 100
                  if eq_stats['total_equations_in_decisions'] > 0 else 0)

    print(f"  텍스트 일치율 (Phase-3): {text_match_pct:.1f}% ({comparison['content_exact_match']}/{comparison['common']})")
    print(f"  수식 Valid율:            {eq_valid_pct:.1f}% ({eq_stats['valid']}/{eq_stats['total_equations_in_decisions']})")
    print(f"  Offset 정합률:           {offset_pct:.1f}% ({eq_stats['offset_exact']}/{eq_stats['total_equations_in_decisions']})")

    # Offset 실패 상세 (처음 5개)
    if eq_stats['offset_failed'] > 0:
        print()
        print("  --- Offset 실패 상세 (처음 5개) ---")
        fail_count = 0
        for d in enriched:
            for eq in d['equations']:
                if eq['char_offset'] < 0:
                    fail_count += 1
                    if fail_count <= 5:
                        print(f"  {d['decision_id']} eq[{eq['position']}]: "
                              f"plain={eq['plain_text'][:40]!r}...")
                        print(f"    content_raw len={len(d['content_raw'])}")
            if fail_count >= 5:
                break


if __name__ == "__main__":
    main()
